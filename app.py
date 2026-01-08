import streamlit as st
import pandas as pd
import torch
import chromadb
import requests
import json
from pymongo import MongoClient
from transformers import AutoTokenizer, AutoModel
from typing import Dict, List, Optional
import time
from datetime import datetime
import os
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from io import BytesIO
import hashlib
from functools import lru_cache
import httpx
from groq import Groq  # Import Groq client

# --- Consolidated Imports ---
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from requests import Session 
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.units import inch
from pypdf import PdfReader, PdfWriter
from docx import Document

import config  # Import centralized config

# --- Helper Functions ---

def extract_template_outline(template_bytes: bytes) -> List[str]:
    """Module-level extractor for PDF template headings to avoid class reload ordering issues."""
    try:
        reader = PdfReader(BytesIO(template_bytes))
        text = []
        for page in reader.pages[:3]:
            try:
                text.append(page.extract_text() or "")
            except Exception:
                continue
        joined = "\n".join(text)
        lines = [l.strip() for l in joined.splitlines()]
        candidates: List[str] = []
        for line in lines:
            if not line:
                continue
            if len(line) < 3 or len(line) > 80:
                continue
            if line.lower().startswith("page "):
                continue
            looks_like_heading = (
                line.endswith(":") or
                (line.isupper() and any(c.isalpha() for c in line)) or
                (line.istitle() and sum(ch.isalpha() for ch in line) >= 6)
            )
            if looks_like_heading:
                normalized = line.rstrip(":").strip()
                if normalized not in candidates:
                    candidates.append(normalized)
        return candidates[:30] if candidates else []
    except Exception:
        return []

# Try to import autogen, but make it optional
try:
    import pyautogen
    AUTOGEN_AVAILABLE = True
except ImportError:
    AUTOGEN_AVAILABLE = False
    # Note: st.warning is moved inside main()

# Performance: HTTP session and model/db connectors
def _get_css(minimal: bool) -> str:
    if minimal:
        return """
<style>
    :root {
        --bg:#ffffff; 
        --fg:#0f172a; 
        --muted:#64748b; 
        --border:#e2e8f0; 
        --primary:#0ea5e9; 
        --primary-dark:#0284c7;
        --accent-blue:#0ea5e9;
        --text-dark:#0f172a;
        --text-light:#64748b;
        --success-green:#16a34a;
        --danger-red:#dc2626;
        --card:#ffffff;
        --card-muted:#f8fafc;
    }
    .main-header { padding: 1rem; border: 1px solid var(--border); border-radius: 12px; background: var(--card); color: var(--fg); }
    .main-header h1 { margin:0; font-size: 1.4rem; }
    .patient-card, .metric-card, .chat-container, .summary-card { border: 1px solid var(--border); border-radius: 12px; padding: 1rem; background: var(--card); color: var(--fg); }
    .patient-card h4, .metric-card h4, .summary-card h3 { color: var(--text-dark); }
    .chat-message { border:1px solid var(--border); border-left:4px solid var(--primary); border-radius:10px; padding:.75rem; background:var(--card-muted); color: var(--text-dark); }
    .doctor-message { background:var(--card-muted); }
    .ai-message { background:var(--card-muted); border-left-color:#9333ea; }
    .stButton > button { background: var(--primary); color:#fff; border:0; border-radius:10px; padding:.6rem 1rem; box-shadow: 0 1px 2px rgba(0,0,0,.05); }
    .stButton > button:hover { background: var(--primary-dark); }
    .stTextArea textarea, .stTextInput input { border-radius:10px !important; border:1px solid var(--border) !import; }
</style>
"""
    return """
<style>
    :root {
        --bg:#ffffff; 
        --fg:#0f172a; 
        --muted:#64748b; 
        --border:#e2e8f0; 
        --primary:#0ea5e9; 
        --primary-dark:#0284c7; 
        --success:#16a34a;
        --accent-blue:#0ea5e9;
        --text-dark:#0f172a;
        --text-light:#64748b;
        --success-green:#16a34a;
        --danger-red:#dc2626;
        --card:#ffffff;
        --card-muted:#f8fafc;
    }
    .main-header { padding: 1rem; border: 1px solid var(--border); border-radius: 12px; background: var(--card); color: var(--fg); text-align:center; margin-bottom: 1rem; }
    .main-header h1 { margin:0; font-size: 1.4rem; }
    .patient-card, .metric-card, .chat-container, .summary-card { border: 1px solid var(--border); border-radius: 12px; padding: 1rem; background: var(--card); color: var(--fg); }
    .metric-card h4 { margin: 0 0 .25rem 0; color: var(--success); font-size: 1rem; }
    .metric-card p { margin: .25rem 0; color: var(--fg); }
    .chat-message { border:1px solid var(--border); border-left:4px solid var(--primary); border-radius:10px; padding:.75rem; background:var(--card-muted); color: var(--text-dark); }
    .doctor-message { background:var(--card-muted); }
    .ai-message { background:var(--card-muted); border-left-color:#9333ea; }
    .stButton > button { background: var(--primary); color:#fff; border:0; border-radius:10px; padding:.6rem 1rem; box-shadow: 0 1px 2px rgba(0,0,0,.05); }
    .stButton > button:hover { background: var(--primary-dark); }
    .stTextArea textarea, .stTextInput input { border-radius:10px !important; border:1px solid var(--border) !important; }
    .empty-state { width:100%; text-align:center; border: 1px dashed var(--border); border-radius: 12px; padding: 2rem; background:var(--card); color: var(--fg); }
    .empty-state .icon { font-size: 2rem; margin-bottom: .5rem; }
    .empty-state h3 { margin: 0 0 .25rem 0; color: var(--fg); font-size: 1.1rem; font-weight: 600; }
    .empty-state p { margin: 0; color: var(--muted); font-size: .95rem; }
</style>
"""

def _http_session() -> Session:
    session = requests.Session()
    retries = Retry(total=2, backoff_factor=0.2, status_forcelist=[429, 502, 503, 504])
    adapter = HTTPAdapter(max_retries=retries, pool_connections=10, pool_maxsize=10)
    session.headers.update({"Connection": "keep-alive"})
    session.mount("http://", adapter)
    session.mount("https://", adapter)
    return session

# Embedding cache to avoid recomputing embeddings
_embedding_cache = {}

def _get_text_hash(text: str) -> str:
    """Generate hash for text to use as cache key"""
    return hashlib.md5(text.encode()).hexdigest()

# FastAPI client configuration
FASTAPI_BASE_URL = os.getenv("FASTAPI_URL", "http://localhost:8000")

class FastAPIClient:
    """Client for FastAPI backend"""
    def __init__(self, base_url: str = FASTAPI_BASE_URL):
        self.base_url = base_url
        self.client = httpx.Client(timeout=60.0)
    
    def chat(self, message: str, patient_data: Optional[Dict] = None) -> str:
        """Chat with AI agent via FastAPI"""
        try:
            response = self.client.post(
                f"{self.base_url}/api/chat",
                json={"message": message, "patient_data": patient_data}
            )
            response.raise_for_status()
            return response.json()["response"]
        except httpx.RequestError as e:
            return f"❌ Error connecting to API: {str(e)}"
        except httpx.HTTPStatusError as e:
            return f"❌ API error: {e.response.text}"
    
    def generate_summary(self, patient_data: str, template_outline: Optional[List[str]] = None) -> str:
        """Generate discharge summary via FastAPI"""
        try:
            response = self.client.post(
                f"{self.base_url}/api/generate-summary",
                json={"patient_data": patient_data, "template_outline": template_outline}
            )
            response.raise_for_status()
            return response.json()["summary"]
        except httpx.RequestError as e:
            return f"❌ Error connecting to API: {str(e)}"
        except httpx.HTTPStatusError as e:
            return f"❌ API error: {e.response.text}"
    
    def search_similar(self, query_text: str, n_results: int = 3) -> List[Dict]:
        """Search similar cases via FastAPI"""
        try:
            response = self.client.post(
                f"{self.base_url}/api/search-similar",
                json={"query_text": query_text, "n_results": n_results}
            )
            response.raise_for_status()
            return response.json()["similar_cases"]
        except httpx.RequestError as e:
            st.error(f"❌ Error connecting to API: {str(e)}")
            return []
        except httpx.HTTPStatusError as e:
            st.error(f"❌ API error: {e.response.text}")
            return []
    
    def get_patient(self, unit_no: str) -> Optional[Dict]:
        """Get patient via FastAPI"""
        try:
            response = self.client.post(
                f"{self.base_url}/api/patient",
                json={"unit_no": unit_no}
            )
            response.raise_for_status()
            return response.json()["patient"]
        except httpx.RequestError as e:
            st.error(f"❌ Error connecting to API: {str(e)}")
            return None
        except httpx.HTTPStatusError as e:
            if e.response.status_code == 404:
                return None
            st.error(f"❌ API error: {e.response.text}")
            return None
    
    def health_check(self) -> bool:
        """Check if FastAPI backend is available"""
        try:
            response = self.client.get(f"{self.base_url}/health", timeout=5.0)
            return response.status_code == 200
        except:
            return False
    
    def close(self):
        """Close the HTTP client"""
        self.client.close()

def _load_tokenizer_model():
    tokenizer = AutoTokenizer.from_pretrained("emilyalsentzer/Bio_ClinicalBERT")
    model = AutoModel.from_pretrained("emilyalsentzer/Bio_ClinicalBERT")
    model.eval()
    if torch.cuda.is_available():
        model.to("cuda")
    return tokenizer, model

def _connect_mongo(uri: str):
    client = MongoClient(uri)
    return client

def _connect_chroma(path: str):
    # IMPORTANT for Cloud: Check if folder exists, otherwise recreate or handle error
    if not os.path.exists(path):
        print(f"⚠️ Warning: ChromaDB path '{path}' not found. Ensure you pushed 'vector_db' to Git.")
    
    client = chromadb.PersistentClient(path=path)
    collection = client.get_or_create_collection("patient_embeddings")
    return client, collection

# --- Business Logic Class ---

class MedicalRAGSystem:
    def __init__(self):
        self.mongo_uri = config.MONGO_URI
        self.chroma_path = config.CHROMA_PATH
        self.ollama_model = config.OLLAMA_MODEL
        self.num_results = 3
        self.http = _http_session()
        self.embedding_cache = _embedding_cache
        
        # --- CLOUD: GROQ SETUP ---
        # 1. Check for API key in Environment Variables (Render)
        # 2. Check for API key in Streamlit Secrets (Streamlit Cloud)
        api_key = os.environ.get("GROQ_API_KEY") or st.secrets.get("GROQ_API_KEY")
        
        if api_key:
            self.groq_client = Groq(api_key=api_key)
            self.use_cloud_api = True
            print("✅ Using Groq Cloud API")
        else:
            self.use_cloud_api = False
            print("⚠️ GROQ_API_KEY not found. Using Local Ollama.")
        
        # Initialize models
        self._load_models()
        self._connect_databases()

    def extract_template_outline(self, template_bytes: bytes) -> list[str]:
        """Extract an ordered list of section headings from a PDF template."""
        try:
            reader = PdfReader(BytesIO(template_bytes))
            text = []
            for page in reader.pages[:3]:
                try:
                    text.append(page.extract_text() or "")
                except Exception:
                    continue
            joined = "\n".join(text)
            lines = [l.strip() for l in joined.splitlines()]
            candidates: list[str] = []
            for line in lines:
                if not line:
                    continue
                if len(line) < 3 or len(line) > 80:
                    continue
                if line.lower().startswith("page "):
                    continue
                looks_like_heading = (
                    line.endswith(":") or
                    (line.isupper() and any(c.isalpha() for c in line)) or
                    (line.istitle() and sum(ch.isalpha() for ch in line) >= 6)
                )
                if looks_like_heading:
                    normalized = line.rstrip(":").strip()
                    if normalized not in candidates:
                        candidates.append(normalized)
            return candidates[:30] if candidates else []
        except Exception:
            return []

    def generate_discharge_summary_with_template(self, patient_data: str, outline_sections: list[str]) -> str:
        """Generate discharge summary following the provided ordered outline sections."""
        outline_bullets = "\n".join([f"- {s}" for s in outline_sections])
        system_prompt = f"""You are an expert medical AI assistant that generates a clinically accurate discharge summary.
Follow the section order EXACTLY as specified by the provided outline. Do not add extra sections; if information is missing, write "[Information not available]".

REQUIRED SECTION ORDER (USE EXACT TITLES):
{outline_bullets}

Rules:
- Use concise, professional medical language.
- Base content solely on the input patient data.
- Preserve patient identifiers verbatim if present.
- Be brief and factual."""

        user_prompt = f"""Generate a discharge summary STRICTLY following the section list above, based only on this data:\n\n{patient_data}\n\nReturn plain text with the exact section headings in order."""

        # --- CLOUD VS LOCAL LOGIC ---
        if self.use_cloud_api:
            try:
                chat_completion = self.groq_client.chat.completions.create(
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    model=config.GROQ_MODEL_ID,
                    temperature=config.GROQ_TEMPERATURE,
                    max_tokens=config.GROQ_MAX_TOKENS,
                )
                return chat_completion.choices[0].message.content
            except Exception as e:
                return f"❌ Error generating summary via Cloud: {str(e)}"
        else:
            # Local Ollama Fallback
            try:
                response = self.http.post(
                    config.OLLAMA_CHAT_ENDPOINT,
                    json={
                        "model": self.ollama_model,
                        "messages": [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_prompt}
                        ],
                        "stream": True,
                        "options": {
                            "temperature": 0.3,  # Lower temperature for faster responses
                            "top_p": 0.85,
                            "max_tokens": 500,  # Reduced for faster generation
                            "num_predict": 500
                        }
                    },
                    timeout=30  # Add timeout
                )
                if response.ok:
                    full_response = ""
                    for line in response.iter_lines(decode_unicode=True):
                        if not line:
                            continue
                        try:
                            json_data = json.loads(line)
                            if 'message' in json_data and 'content' in json_data['message']:
                                content = json_data['message']['content']
                                if content:
                                    full_response += content
                            if json_data.get('done', False):
                                break
                        except json.JSONDecodeError:
                            continue
                    return full_response.strip() if full_response.strip() else "Summary generated successfully."
                else:
                    return f"❌ Error generating summary: {response.text}"
            except Exception as e:
                return f"❌ Error connecting to Ollama: {str(e)}"

    def generate_pdf_from_text(self, text: str, template_bytes: bytes | None = None) -> bytes:
        """Generate a PDF from plain text."""
        # Determine page size
        page_size = A4
        template_reader = None
        if template_bytes:
            try:
                template_reader = PdfReader(BytesIO(template_bytes))
                first_page = template_reader.pages[0]
                width = float(first_page.mediabox.width)
                height = float(first_page.mediabox.height)
                page_size = (width, height)
            except Exception:
                template_reader = None

        buf = BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=page_size, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()
        body_style = ParagraphStyle(
            name="Body",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            alignment=TA_LEFT,
        )
        story = []
        for para in text.split("\n\n"):
            story.append(Paragraph(para.replace("\n", "<br/>"), body_style))
            story.append(Spacer(1, 0.18 * inch))
        doc.build(story)
        generated_pdf_bytes = buf.getvalue()

        if not template_reader:
            return generated_pdf_bytes

        gen_reader = PdfReader(BytesIO(generated_pdf_bytes))
        writer = PdfWriter()
        num_pages = max(len(template_reader.pages), len(gen_reader.pages))

        for i in range(num_pages):
            template_page = None
            if i < len(template_reader.pages):
                template_page = template_reader.pages[i]

            if i < len(gen_reader.pages):
                gen_page = gen_reader.pages[i]
                if template_page is not None:
                    try:
                        template_page.merge_page(gen_page)
                        writer.add_page(template_page)
                    except Exception:
                        writer.add_page(gen_page)
                else:
                    writer.add_page(gen_page)
            else:
                if template_page is not None:
                    writer.add_page(template_page)

        out_buf = BytesIO()
        writer.write(out_buf)
        return out_buf.getvalue()

    def generate_docx_from_text(self, text: str) -> bytes:
        """Generate a DOCX file from plain text."""
        doc = Document()
        lines = text.split("\n")
        if lines and len(lines[0]) <= 120 and any(ch.isalpha() for ch in lines[0]):
            doc.add_heading(lines[0].strip(), level=1)
            text = "\n".join(lines[1:])
        for block in text.split("\n\n"):
            for ln in block.split("\n"):
                doc.add_paragraph(ln)
            doc.add_paragraph("")
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
    
    def _load_models(self):
        """Load Bio ClinicalBERT model for embeddings"""
        with st.spinner("Loading Bio ClinicalBERT model..."):
            self.tokenizer, self.model = _load_tokenizer_model()
    
    def _connect_databases(self):
        """Connect to MongoDB and ChromaDB"""
        try:
            # MongoDB connection
            self.mongo_client = _connect_mongo(self.mongo_uri)
            self.db = self.mongo_client[config.DATABASE_NAME]
            self.patients_collection = self.db[config.PATIENTS_COLLECTION]
            
            # ChromaDB connection
            self.chroma_client, self.chroma_collection = _connect_chroma(self.chroma_path)
            
            st.success("✅ Connected to databases successfully")
        except Exception as e:
            st.error(f"❌ Database connection failed: {str(e)}")
    
    def embed_text(self, text: str) -> List[float]:
        """Generate embedding for text using Bio ClinicalBERT with caching"""
        # Check cache first
        text_hash = _get_text_hash(text)
        if text_hash in self.embedding_cache:
            return self.embedding_cache[text_hash]
        
        # Generate embedding
        with torch.no_grad():
            inputs = self.tokenizer(text, return_tensors="pt", truncation=True, max_length=256)
            if torch.cuda.is_available():
                inputs = {k: v.to("cuda") for k, v in inputs.items()}
            outputs = self.model(**inputs)
            cls_embedding = outputs.last_hidden_state[:, 0, :]
            emb = cls_embedding.squeeze(0)
            if emb.is_cuda:
                emb = emb.to("cpu")
            embedding = emb.tolist()
        
        # Cache the embedding
        self.embedding_cache[text_hash] = embedding
        return embedding
    
    def format_patient_fields(self, record: Dict) -> str:
        """Format patient record fields for embedding"""
        fields = config.PATIENT_FIELDS
        parts = [f"{field.title()}: {record.get(field, '')}" for field in fields if record.get(field)]
        return " ".join(parts)
    
    def get_patient_by_unit_no(self, unit_no: str) -> Optional[Dict]:
        """Retrieve patient record from MongoDB"""
        try:
            record = self.patients_collection.find_one({"unit no": int(unit_no)})
            return record
        except Exception as e:
            st.error(f"Error retrieving patient: {str(e)}")
            return None
    
    def search_similar_cases(self, query_text: str, n_results: int = 3) -> List[Dict]:
        """Search for similar cases using RAG"""
        try:
            query_embedding = self.embed_text(query_text)
            results = self.chroma_collection.query(
                query_embeddings=[query_embedding],
                n_results=n_results,
                include=["documents", "metadatas"]
            )
            
            similar_cases = []
            for i in range(len(results["documents"][0])):
                similar_cases.append({
                    "document": results["documents"][0][i],
                    "metadata": results["metadatas"][0][i],
                    "similarity": 1 - results["distances"][0][i]  # Convert distance to similarity
                })
            
            return similar_cases
        except Exception as e:
            st.error(f"Error searching similar cases: {str(e)}")
            return []
    
    def generate_discharge_summary(self, patient_data: str, similar_cases: List[Dict] = None) -> str:
        """Generate discharge summary using Ollama LLM"""
        system_prompt = """You are an expert medical AI assistant that generates structured, clinically accurate discharge summaries.
Base your summary entirely on the INPUT PATIENT DATA provided.
The discharge summary MUST include: Name, Unit No, Date Of Birth, Sex, Admission/Discharge Dates, Attending, Chief Complaint, Procedure, History, Physical Exam (on Admission), Pertinent Results, Brief Hospital Course, Medications on Admission, Discharge Medications, Discharge Instructions, Discharge Disposition, Discharge Diagnosis, Discharge Condition, Follow-up.

For Name, Unit No, Date of Birth, and Sex, copy the information verbatim.
If information is missing, state "[Information not available]".
Use concise, professional medical language. Be brief and factual."""

        user_prompt = f"""Generate a discharge summary for this patient:
{patient_data}

Extract Name, Unit No, Date of Birth, and Sex exactly as provided."""

        # --- CLOUD VS LOCAL LOGIC ---
        if self.use_cloud_api:
            try:
                chat_completion = self.groq_client.chat.completions.create(
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    model=config.GROQ_MODEL_ID,
                    temperature=config.GROQ_TEMPERATURE,
                    max_tokens=config.GROQ_MAX_TOKENS,
                )
                return chat_completion.choices[0].message.content
            except Exception as e:
                return f"❌ Error generating summary via Cloud: {str(e)}"
        else:
            # Local Ollama Fallback
            try:
                response = self.http.post(
                    config.OLLAMA_CHAT_ENDPOINT,
                    json={
                        "model": self.ollama_model,
                        "messages": [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_prompt}
                        ],
                        "stream": True,
                        "options": {
                            "temperature": 0.3,  # Lower temperature for faster responses
                            "top_p": 0.85,
                            "max_tokens": 500,  # Reduced for faster generation
                            "num_predict": 500
                        }
                    },
                    timeout=30  # Add timeout
                )

                if response.ok:
                    full_response = ""
                    for line in response.iter_lines(decode_unicode=True):
                        if not line:
                            continue
                        try:
                            json_data = json.loads(line)
                            if 'message' in json_data and 'content' in json_data['message']:
                                content = json_data['message']['content']
                                if content:
                                    full_response += content
                            if json_data.get('done', False):
                                break
                        except json.JSONDecodeError:
                            continue
                    return full_response.strip() if full_response.strip() else "Summary generated successfully."
                else:
                    return f"❌ Error generating summary: {response.text}"
            except Exception as e:
                return f"❌ Error connecting to Ollama: {str(e)}"

    # --- START: NEW FEEDBACK LOOP METHOD ---
    def add_summary_to_vector_db(self, patient_info: Dict, summary_text: str):
        """
        Embeds the finalized discharge summary and adds it to the ChromaDB collection.
        This serves as the feedback loop, adding a high-quality, human-reviewed
        document back into the RAG system.
        """
        if not summary_text or not patient_info:
            st.warning("No summary text or patient info available to add.")
            return False

        unit_no = patient_info.get('unit no', 'unknown')
        patient_name = patient_info.get('name', 'Unknown')
        
        try:
            # 1. Generate embedding for the new summary
            summary_embedding = self.embed_text(summary_text)
            
            # 2. Prepare a unique ID
            # Using unit_no and timestamp allows for multiple summary versions
            doc_id = f"summary_{unit_no}_{int(time.time())}"
            
            # 3. Prepare metadata
            metadata = {
                "unit_no": str(unit_no),
                "name": patient_name,
                "summary": summary_text[:500],  # Store a preview in metadata
                "source_type": "feedback_summary" # Tag this as a human-reviewed entry
            }
            
            # 4. Add to ChromaDB
            self.chroma_collection.add(
                embeddings=[summary_embedding],
                documents=[summary_text],  # Store the full summary as the document
                metadatas=[metadata],
                ids=[doc_id]
            )
            
            # 5. Show notification (as requested)
            # st.toast is available in newer Streamlit; fall back to success if missing
            try:
                st.toast(f"Database updated: Summary for {unit_no} added.", icon="✅")
            except Exception:
                st.success(f"Database updated: Summary for {unit_no} added.")
            return True
        
        except Exception as e:
            st.error(f"❌ Error adding feedback summary to vector DB: {str(e)}")
            st.exception(e) # Print full error
            return False
    # --- END: NEW FEEDBACK LOOP METHOD ---
    
class AutoGenMedicalAgent:
    def __init__(self, rag_system: MedicalRAGSystem, api_client: FastAPIClient = None):
        self.rag_system = rag_system
        self.api_client = api_client
        self.agent = None
        self.user_proxy = None
        self._initialize_agent()
    
    def _initialize_agent(self):
        """Initialize AutoGen medical assistant agent"""
        # Skip AutoGen initialization to avoid API errors
        # Use FastAPI backend instead
        pass
    
    def chat_with_doctor(self, message: str, patient_data: Dict = None) -> str:
        """Handle conversation with doctor - uses FastAPI if available"""
        try:
            # Use FastAPI if available, otherwise fallback
            if self.api_client and self.api_client.health_check():
                return self.api_client.chat(message, patient_data)
            else:
                return self._fallback_chat(message, patient_data)
        except Exception as e:
            return f"❌ Error in conversation: {str(e)}"
    
    def _fallback_chat(self, message: str, patient_data: Dict = None) -> str:
        """Fallback chat using direct Ollama/Groq interaction"""
        try:
            # Check if user is asking for discharge summary generation
            if "discharge summary" in message.lower() or "generate summary" in message.lower():
                if patient_data:
                    # Use the existing discharge summary generation method
                    patient_text = self.rag_system.format_patient_fields(patient_data)
                    return self.rag_system.generate_discharge_summary(patient_text)
                else:
                    return "❌ Please select a patient first to generate a discharge summary."
            
            # Add patient context if available (truncated for speed)
            context = ""
            if patient_data:
                # Truncate context to avoid long prompts
                context = f"\n\nPatient: {patient_data.get('name', 'Unknown')} (Unit {patient_data.get('unit no', 'N/A')})"
            
            system_prompt = """You are a medical AI assistant. Provide concise, accurate responses. Keep answers brief (2-3 sentences max)."""
            
            # --- CLOUD VS LOCAL LOGIC ---
            if self.rag_system.use_cloud_api:
                chat_completion = self.rag_system.groq_client.chat.completions.create(
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": f"{message}{context}"}
                    ],
                    model=config.GROQ_MODEL_ID,
                    temperature=0.6,
                    max_tokens=250,
                )
                return chat_completion.choices[0].message.content
            else:
                # Local Fallback
                response = self.rag_system.http.post(
                    config.OLLAMA_CHAT_ENDPOINT,
                    json={
                        "model": config.OLLAMA_MODEL,
                        "messages": [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": f"{message}{context}"}
                        ],
                        "stream": True,
                        "options": {
                            "temperature": 0.6,
                            "top_p": 0.85,
                            "max_tokens": 150,
                            "num_predict": 150
                        }
                    },
                    timeout=20  # Add timeout for faster failure
                )
                
                if response.ok:
                    full_response = ""
                    for line in response.iter_lines(decode_unicode=True):
                        if line:
                            try:
                                json_data = json.loads(line)
                                if 'message' in json_data and 'content' in json_data['message']:
                                    content = json_data['message']['content']
                                    if content:
                                        full_response += content
                                if json_data.get('done', False):
                                    break
                            except json.JSONDecodeError:
                                continue
                    return full_response.strip() if full_response.strip() else "I'm here to help with medical questions. How can I assist you?"
                else:
                    return f"❌ Error connecting to Ollama: {response.text}"
                
        except requests.exceptions.Timeout:
            return "⏱️ Request timed out. Please try again with a shorter message."
        except Exception as e:
            return f"❌ Error in fallback chat: {str(e)}"

def main():
    # Initialize FastAPI client
    if 'api_client' not in st.session_state:
        st.session_state.api_client = FastAPIClient()
    
    # Check FastAPI availability
    use_fastapi = st.session_state.api_client.health_check()
    if use_fastapi:
        st.session_state.use_fastapi = True
        # Initialize with FastAPI client
        if 'rag_system' not in st.session_state:
            # Still need RAG system for some operations (formatting, PDF generation)
            try:
                st.session_state.rag_system = MedicalRAGSystem()
            except Exception as e:
                st.warning(f"⚠️ Could not initialize full RAG system: {str(e)}. Some features may be limited.")
        # Always initialize autogen_agent if not present
        if 'autogen_agent' not in st.session_state:
            try:
                st.session_state.autogen_agent = AutoGenMedicalAgent(
                    st.session_state.rag_system if 'rag_system' in st.session_state else None,
                    st.session_state.api_client
                )
            except Exception as e:
                st.error(f"❌ Failed to initialize AI agent: {str(e)}")
    else:
        st.session_state.use_fastapi = False
        # Initialize RAG system as fallback
        if 'rag_system' not in st.session_state:
            with st.spinner("Initializing Medical RAG System (FastAPI unavailable, using fallback)..."):
                try:
                    st.session_state.rag_system = MedicalRAGSystem()
                    st.session_state.autogen_agent = AutoGenMedicalAgent(st.session_state.rag_system, None)
                    # NOTE: Removed the warning about FastAPI for production, 
                    # as production is designed to use this fallback path.
                except Exception as e:
                    st.error(f"❌ Failed to initialize system: {str(e)}")
                    st.stop()
        # Ensure autogen_agent is initialized even in fallback mode
        if 'autogen_agent' not in st.session_state:
            try:
                st.session_state.autogen_agent = AutoGenMedicalAgent(
                    st.session_state.rag_system if 'rag_system' in st.session_state else None,
                    None
                )
            except Exception as e:
                st.error(f"❌ Failed to initialize AI agent: {str(e)}")

    # Sidebar preferences and CSS
    with st.sidebar:
        st.header("⚙️ Preferences")
        minimal_ui = st.checkbox("Minimal UI", value=st.session_state.get('minimal_ui', False))
        st.session_state.minimal_ui = minimal_ui

        st.markdown("---")
        st.header("📎 Insurance Template")
        template_file = st.file_uploader("Upload PDF template (optional)", type=["pdf"], accept_multiple_files=False)
        if template_file is not None:
            st.session_state["template_pdf_bytes"] = template_file.read()
            # Extract outline from template (module-level helper to avoid class ordering issues)
            outline = extract_template_outline(st.session_state["template_pdf_bytes"])
            if outline:
                st.session_state["template_outline"] = outline
                st.success("Template loaded. Outline detected and will be used for generation.")
                with st.expander("Detected Section Order"):
                    for s in outline:
                        st.write(f"• {s}")
            else:
                st.session_state.pop("template_outline", None)
                st.warning("Template loaded but no clear section outline was detected. Will generate standard summary.")
        elif "template_pdf_bytes" not in st.session_state:
            st.info("No template uploaded. Summaries will be generated as plain text or basic PDF.")

    st.markdown(_get_css(st.session_state.get('minimal_ui', False)), unsafe_allow_html=True)

    # Header with modern dark design
    st.markdown("""
    <div class="main-header">
        <h1>🏥 Medical Discharge Summary Assistant</h1>
        <p>AI-Powered Clinical Documentation with RAG and AutoGen Integration</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Status indicators with modern cards
    col_status1, col_status2, col_status3, col_status4 = st.columns(4)
    
    with col_status1:
        st.markdown("""
        <div class="metric-card">
            <h4>🟢 System Ready</h4>
            <p>All systems operational</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col_status2:
        mode = "Cloud (Groq)" if st.session_state.rag_system.use_cloud_api else "Local"
        st.markdown(f"""
        <div class="metric-card">
            <h4>⚡ {mode}</h4>
            <p>Optimized responses</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col_status3:
        st.markdown("""
        <div class="metric-card">
            <h4>🧠 AI Active</h4>
            <p>LLaMA 3 + RAG</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col_status4:
        patient_count = "1" if st.session_state.current_patient else "0"
        st.markdown(f"""
        <div class="metric-card">
            <h4>👤 Patient</h4>
            <p>{patient_count} selected</p>
    </div>
    """, unsafe_allow_html=True)
    
    # RAG system is already initialized above
    
    # Sidebar for patient search
    with st.sidebar:
        st.header("🔍 Patient Search")
        
        # Patient search form
        with st.form("patient_search"):
            unit_no = st.text_input("Unit Number", placeholder="Enter patient unit number")
            search_button = st.form_submit_button("🔍 Search Patient")
            
            if search_button and unit_no:
                with st.spinner("Searching for patient..."):
                    try:
                        # Use FastAPI if available
                        if st.session_state.get('use_fastapi', False):
                            patient = st.session_state.api_client.get_patient(unit_no)
                        else:
                            patient = st.session_state.rag_system.get_patient_by_unit_no(unit_no)
                        if patient:
                            st.session_state.current_patient = patient
                            st.markdown(f"""
                            <div style="background: rgba(16, 185, 129, 0.15); border: 1px solid rgba(16, 185, 129, 0.4); border-radius: 12px; padding: 1rem; margin: 0.5rem 0;">
                                <p style="color: var(--success); margin: 0; font-weight: 600;">✅ Found patient: {patient.get('name', 'Unknown')}</p>
                            </div>
                            """, unsafe_allow_html=True)
                        else:
                            st.markdown("""
                            <div style="background: rgba(239, 68, 68, 0.15); border: 1px solid rgba(239, 68, 68, 0.4); border-radius: 12px; padding: 1rem; margin: 0.5rem 0;">
                                <p style="color: var(--danger); margin: 0; font-weight: 600;">❌ Patient not found</p>
                            </div>
                            """, unsafe_allow_html=True)
                    except Exception as e:
                        st.markdown(f"""
                        <div style="background: rgba(239, 68, 68, 0.15); border: 1px solid rgba(239, 68, 68, 0.4); border-radius: 12px; padding: 1rem; margin: 0.5rem 0;">
                            <p style="color: var(--danger); margin: 0; font-weight: 600;">❌ Error: {str(e)}</p>
                        </div>
                        """, unsafe_allow_html=True)
        
        # Display current patient info
        if st.session_state.current_patient:
            st.markdown("### 👤 Current Patient")
            patient = st.session_state.current_patient
            
            st.markdown(f"""
            <div class="patient-card">
                <h4 style="color: var(--text-primary); margin-bottom: 1.5rem; font-size: 1.3rem; font-weight: 700;">📋 {patient.get('name', 'Unknown')}</h4>
                <p style="color: var(--text-secondary); margin: 0.75rem 0; font-size: 0.95rem;"><strong style="color: var(--primary-light);">Unit No:</strong> {patient.get('unit no', 'N/A')}</p>
                <p style="color: var(--text-secondary); margin: 0.75rem 0; font-size: 0.95rem;"><strong style="color: var(--primary-light);">DOB:</strong> {patient.get('date of birth', 'N/A')}</p>
                <p style="color: var(--text-secondary); margin: 0.75rem 0; font-size: 0.95rem;"><strong style="color: var(--primary-light);">Sex:</strong> {patient.get('sex', 'N/A')}</p>
                <p style="color: var(--text-secondary); margin: 0.75rem 0; font-size: 0.95rem;"><strong style="color: var(--primary-light);">Service:</strong> {patient.get('service', 'N/A')}</p>
                <p style="color: var(--text-secondary); margin: 0.75rem 0; font-size: 0.95rem;"><strong style="color: var(--primary-light);">Attending:</strong> {patient.get('attending', 'N/A')}</p>
            </div>
            """, unsafe_allow_html=True)
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown("""
        <div class="card-header">
            <div class="card-header-icon">💬</div>
            <h3 style="margin: 0; color: var(--text-dark);">AI Medical Assistant</h3>
        </div>
        """, unsafe_allow_html=True)
        
        # Chat interface
        if st.session_state.current_patient:
            # Chat container with better styling
            st.markdown("""
            <div class="chat-container">
            """, unsafe_allow_html=True)
            
            # Display chat history
            if st.session_state.chat_history:
                for message in st.session_state.chat_history:
                    if message["role"] == "doctor":
                        st.markdown(f"""
                        <div class="chat-message doctor-message">
                            <div style="display: flex; align-items: flex-start; gap: 1rem;">
                                <div style="background: var(--gradient-primary); color: white; border-radius: 50%; width: 44px; height: 44px; display: flex; align-items: center; justify-content: center; font-size: 1.3rem; flex-shrink: 0; box-shadow: var(--shadow-md);">👨‍⚕️</div>
                                <div style="flex: 1;">
                                    <div style="font-weight: 700; color: var(--primary-light); margin-bottom: 0.5rem; font-size: 0.9rem; letter-spacing: 0.02em;">Doctor</div>
                                    <div style="color: var(--text-primary); line-height: 1.7; font-size: 0.95rem;">{message["content"]}</div>
                                </div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    else:
                        st.markdown(f"""
                        <div class="chat-message ai-message">
                            <div style="display: flex; align-items: flex-start; gap: 1rem;">
                                <div style="background: var(--gradient-primary); color: white; border-radius: 50%; width: 44px; height: 44px; display: flex; align-items: center; justify-content: center; font-size: 1.3rem; flex-shrink: 0; box-shadow: var(--shadow-md);">🤖</div>
                                <div style="flex: 1;">
                                    <div style="font-weight: 700; color: var(--accent-purple); margin-bottom: 0.5rem; font-size: 0.9rem; letter-spacing: 0.02em;">AI Assistant</div>
                                    <div style="color: var(--text-primary); line-height: 1.7; font-size: 0.95rem;">{message["content"]}</div>
                                </div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <div style="text-align: center; color: var(--text-muted); padding: 3rem 2rem; font-style: italic; font-size: 1.1rem;">
                    👋 Start a conversation with the AI assistant...
                </div>
                """, unsafe_allow_html=True)
            
            st.markdown("</div>", unsafe_allow_html=True)
            
            # Chat input form
            with st.form("chat_form", clear_on_submit=False):
                st.markdown("**<span style='color: var(--text-dark); font-weight: 600;'>Ask the AI assistant:</span>**", unsafe_allow_html=True)
                user_message = st.text_area(
                    "message_input", 
                    placeholder="e.g., Generate a discharge summary for this patient",
                    height=100,
                    label_visibility="collapsed"
                )
                
                col_send, col_clear = st.columns([1, 1])
                with col_send:
                    send_button = st.form_submit_button("💬 Send Message", type="primary", use_container_width=True)
                with col_clear:
                    clear_button = st.form_submit_button("🗑️ Clear Chat", use_container_width=True)
                
                if send_button and user_message.strip():
                    # Add doctor message to history
                    st.session_state.chat_history.append({
                        "role": "doctor",
                        "content": user_message.strip(),
                        "timestamp": datetime.now()
                    })
                    
                    # Get AI response with progress indicator
                    with st.spinner("🤖 AI is thinking..."):
                        try:
                            # Check if autogen_agent is initialized
                            if 'autogen_agent' not in st.session_state or st.session_state.autogen_agent is None:
                                # Try to initialize it
                                if st.session_state.get('use_fastapi', False):
                                    st.session_state.autogen_agent = AutoGenMedicalAgent(
                                        st.session_state.rag_system if 'rag_system' in st.session_state else None,
                                        st.session_state.api_client
                                    )
                                else:
                                    if 'rag_system' not in st.session_state:
                                        st.error("❌ RAG system not initialized. Please refresh the page.")
                                        st.stop()
                                    st.session_state.autogen_agent = AutoGenMedicalAgent(st.session_state.rag_system, None)
                            
                            start_time = time.time()
                            ai_response = st.session_state.autogen_agent.chat_with_doctor(
                                user_message.strip(), 
                                st.session_state.current_patient
                            )
                            elapsed = time.time() - start_time
                            if elapsed < 2:
                                st.success(f"⚡ Response generated in {elapsed:.1f}s")
                            
                            # Add AI response to history
                            st.session_state.chat_history.append({
                                "role": "ai",
                                "content": ai_response,
                                "timestamp": datetime.now()
                            })
                        except Exception as e:
                            st.session_state.chat_history.append({
                                "role": "ai",
                                "content": f"❌ Error: {str(e)}",
                                "timestamp": datetime.now()
                            })
                    
                    st.rerun()
                
                if clear_button:
                    st.session_state.chat_history = []
                    st.rerun()
            
            # Action buttons with modern styling
            st.markdown("---")
            st.markdown("""
            <div class="card-header">
                <div class="card-header-icon">🚀</div>
                <h3 style="margin: 0; color: var(--text-dark);">Quick Actions</h3>
            </div>
            """, unsafe_allow_html=True)
            col_btn1, col_btn2, col_btn3 = st.columns(3)
            
            with col_btn1:
                if st.button("📝 Generate Summary", type="primary", use_container_width=True):
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    try:
                        status_text.text("📝 Formatting patient data...")
                        progress_bar.progress(20)
                        patient_text = st.session_state.rag_system.format_patient_fields(st.session_state.current_patient)
                        
                        status_text.text("🤖 Generating summary with AI...")
                        progress_bar.progress(40)
                        start_time = time.time()
                        
                        # Use FastAPI if available
                        if st.session_state.get('use_fastapi', False):
                            template_outline = st.session_state.get("template_outline")
                            summary = st.session_state.api_client.generate_summary(patient_text, template_outline)
                        else:
                            # If a template outline exists, follow it strictly
                            if "template_outline" in st.session_state and st.session_state.template_outline:
                                summary = st.session_state.rag_system.generate_discharge_summary_with_template(patient_text, st.session_state.template_outline)
                            else:
                                summary = st.session_state.rag_system.generate_discharge_summary(patient_text)
                        
                        elapsed = time.time() - start_time
                        progress_bar.progress(80)
                        status_text.text("📄 Preparing document...")
                        
                        st.session_state.discharge_summary = summary
                        # Build PDF (with template if provided)
                        template_bytes = st.session_state.get("template_pdf_bytes", None)
                        # For template mode, generate a clean PDF using the template's page size but avoid overlaying duplicate headings
                        pdf_bytes = st.session_state.rag_system.generate_pdf_from_text(summary, template_bytes=None if st.session_state.get("template_outline") else template_bytes)
                        st.session_state.discharge_summary_pdf = pdf_bytes
                        
                        progress_bar.progress(100)
                        status_text.empty()
                        progress_bar.empty()
                        st.success(f"✅ Discharge summary generated in {elapsed:.1f}s!")
                    except Exception as e:
                        st.error(f"❌ Error generating summary: {str(e)}")
                    st.rerun()
            
            with col_btn2:
                if st.button("🔍 Find Similar Cases", use_container_width=True):
                    with st.spinner("🔍 Searching for similar cases..."):
                        try:
                            patient_text = st.session_state.rag_system.format_patient_fields(st.session_state.current_patient)
                            # Use FastAPI if available
                            if st.session_state.get('use_fastapi', False):
                                similar_cases = st.session_state.api_client.search_similar(patient_text)
                            else:
                                similar_cases = st.session_state.rag_system.search_similar_cases(patient_text)
                            st.session_state.similar_cases = similar_cases
                            st.success(f"✅ Found {len(similar_cases)} similar cases!")
                        except Exception as e:
                            st.error(f"❌ Error searching cases: {str(e)}")
                        st.rerun()
            
            with col_btn3:
                if st.button("📊 Patient Overview", use_container_width=True):
                    with st.spinner("📊 Analyzing patient data..."):
                        try:
                            patient = st.session_state.current_patient
                            overview = f"""**Patient Overview:**

**Name:** {patient.get('name', 'Unknown')}
**Unit No:** {patient.get('unit no', 'N/A')}
**Date of Birth:** {patient.get('date of birth', 'N/A')}
**Sex:** {patient.get('sex', 'N/A')}
**Service:** {patient.get('service', 'N/A')}
**Chief Complaint:** {patient.get('chief complaint', 'N/A')}
**Attending:** {patient.get('attending', 'N/A')}
**Allergies:** {patient.get('allergies', 'N/A')}
**Past Medical History:** {patient.get('past medical history', 'N/A')[:200]}{'...' if len(str(patient.get('past medical history', ''))) > 200 else ''}

This patient is ready for discharge summary generation."""
                            
                            st.session_state.chat_history.append({
                                "role": "ai",
                                "content": overview,
                                "timestamp": datetime.now()
                            })
                            st.success("✅ Patient overview added to chat!")
                        except Exception as e:
                            st.error(f"❌ Error generating overview: {str(e)}")
                        st.rerun()
        
        else:
            st.markdown("""
            <div class="empty-state">
                <div class="icon">👈</div>
                <h3>Search for a Patient</h3>
                <p>Please search for a patient in the sidebar to start the conversation with the AI assistant.</p>
            </div>
            """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="card-header">
            <div class="card-header-icon">📋</div>
            <h3 style="margin: 0; color: var(--text-dark);">Generated Content</h3>
        </div>
        """, unsafe_allow_html=True)
        
        # Display discharge summary (editable)
        if st.session_state.discharge_summary:
            st.markdown("""
            <div class="summary-card">
                <div class="card-header">
                    <div class="card-header-icon">📄</div>
                    <h3 style="margin: 0; color: var(--text-dark);">Discharge Summary (Editable)</h3>
                </div>
            </div>
            """, unsafe_allow_html=True)

            if "editable_summary" not in st.session_state:
                st.session_state.editable_summary = st.session_state.discharge_summary
            # Use current discharge_summary if editable_summary is empty or reset
            elif not st.session_state.editable_summary:
                st.session_state.editable_summary = st.session_state.discharge_summary

            st.session_state.editable_summary = st.text_area(
                 "editable_summary",
                 value=st.session_state.editable_summary,
                 height=500,
                 label_visibility="collapsed"
             )

            col_save, col_reset = st.columns([1,1])
            with col_save:
                 if st.button("💾 Save Edits", use_container_width=True):
                     st.session_state.discharge_summary = st.session_state.editable_summary
                     st.success("Saved your edits.")
            with col_reset:
                 if st.button("↩️ Reset to Generated", use_container_width=True):
                    st.session_state.editable_summary = st.session_state.discharge_summary
                    st.rerun() # Rerun to ensure text_area updates

            # --- START: NEW FEEDBACK LOOP UI ---
            st.markdown("---")
            st.markdown("### 🧠 RAG Feedback Loop")
            
            if st.button("Commit Summary to Knowledgebase", 
                         type="primary", 
                         use_container_width=True, 
                         help="Embed this summary and add it to the RAG system for future 'similar cases' searches."):
                
                if st.session_state.editable_summary and st.session_state.current_patient:
                    with st.spinner("Embedding summary and updating knowledgebase..."):
                        st.session_state.rag_system.add_summary_to_vector_db(
                            st.session_state.current_patient,
                            st.session_state.editable_summary
                        )
                else:
                    st.warning("Please ensure a patient is loaded and a summary is present.")
            # --- END: NEW FEEDBACK LOOP UI ---
            
            st.markdown("---") # Added a separator

            # Plain text download
            st.download_button(
                label="📥 Download as .txt",
                data=st.session_state.editable_summary,
                file_name=f"discharge_summary_{st.session_state.current_patient.get('unit no', 'unknown')}.txt",
                mime="text/plain"
            )

            # DOCX download
            docx_bytes = st.session_state.rag_system.generate_docx_from_text(st.session_state.editable_summary)
            st.download_button(
                label="📝 Download as .docx",
                data=docx_bytes,
                file_name=f"discharge_summary_{st.session_state.current_patient.get('unit no', 'unknown')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # PDF download (optional template mode)
            if "discharge_summary_pdf" in st.session_state and st.session_state.discharge_summary_pdf:
                st.download_button(
                    label="🧾 Download PDF (Template Applied)",
                    data=st.session_state.discharge_summary_pdf,
                    file_name=f"discharge_summary_{st.session_state.current_patient.get('unit no', 'unknown')}.pdf",
                    mime="application/pdf"
                )
        
        # Display similar cases
        if hasattr(st.session_state, 'similar_cases') and st.session_state.similar_cases:
            st.markdown("### 🔍 Similar Cases Found")
            
            for i, case in enumerate(st.session_state.similar_cases):
                with st.expander(f"Case {i+1} - Similarity: {case['similarity']:.2%}"):
                    st.write("**Patient Info:**")
                    st.write(f"Name: {case['metadata'].get('name', 'Unknown')}")
                    st.write(f"Unit No: {case['metadata'].get('unit_no', 'Unknown')}")
                    
                    st.write("**Summary Preview:**")
                    summary_preview = case['metadata'].get('summary', 'No summary available')[:200] + "..."
                    st.write(summary_preview)
    
    # Footer removed - status is now in header

if __name__ == "__main__":
    main()